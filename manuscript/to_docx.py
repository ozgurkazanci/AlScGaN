"""Render MANUSCRIPT.md as a formatted Word document.

Written against python-docx directly because pandoc, LibreOffice and node are
not available on this machine. It handles exactly the Markdown this manuscript
uses - headings, bold and italic spans, bullet lists, pipe tables, images and
horizontal rules - rather than pretending to be a general converter.
"""

import os
import pathlib
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = pathlib.Path(__file__).parent
SRC = HERE / "MANUSCRIPT.md"
OUT = HERE / "MANUSCRIPT.docx"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)
INK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x52, 0x51, 0x4E)

# A4 with 2.5 cm margins leaves this much for figures and tables
CONTENT_WIDTH = Cm(21.0 - 2 * 2.5)


# --- document setup -----------------------------------------------------------
def setup(doc):
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        for attr in ("left_margin", "right_margin"):
            setattr(s, attr, Cm(2.5))
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = INK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), BODY_FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.4

    for name, size, bold, before, after in (
            ("Heading 1", 15, True, 18, 8),
            ("Heading 2", 12.5, True, 14, 6),
            ("Heading 3", 11, True, 10, 4)):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = INK
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_page_numbers(doc):
    """Field codes for 'page N' in the footer."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for tag, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if tag:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), tag)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = f" {text} "
            run._r.append(instr)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.font.name = BODY_FONT


# --- inline formatting --------------------------------------------------------
INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+`)")


def add_runs(par, text, base_italic=False, size=None, color=None):
    for piece in INLINE.split(text):
        if not piece:
            continue
        bold = italic = mono = False
        if piece.startswith("**") and piece.endswith("**"):
            piece, bold = piece[2:-2], True
        elif piece.startswith("*") and piece.endswith("*"):
            piece, italic = piece[1:-1], True
        elif piece.startswith("`") and piece.endswith("`"):
            piece, mono = piece[1:-1], True
        run = par.add_run(piece)
        run.bold = bold
        run.italic = italic or base_italic
        run.font.name = "Consolas" if mono else BODY_FONT
        if size:
            run.font.size = size
        if color is not None:
            run.font.color.rgb = color
    return par


# --- block elements -----------------------------------------------------------
def add_callout(doc, text):
    """A shaded paragraph, used for the simulated-data disclaimer."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0EFEC")
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "EB6834")
    left.set(qn("w:space"), "6")
    pbdr.append(left)
    ppr = p._p.get_or_add_pPr()
    ppr.append(shd)
    ppr.append(pbdr)
    add_runs(p, text)
    return p


def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "C3C2B7")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1 + len(body), cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    col_w = int(CONTENT_WIDTH.emu / len(header))
    for i, cell in enumerate(t.rows[0].cells):
        cell.width = col_w
        cell.paragraphs[0].text = ""
        add_runs(cell.paragraphs[0], f"**{header[i]}**", size=Pt(9.5))
        shade = OxmlElement("w:shd")
        shade.set(qn("w:val"), "clear")
        shade.set(qn("w:fill"), "F0EFEC")
        cell._tc.get_or_add_tcPr().append(shade)
    for r, row in enumerate(body, start=1):
        for c, val in enumerate(row):
            if c >= len(header):
                continue
            cell = t.rows[r].cells[c]
            cell.width = col_w
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], val, size=Pt(9.5))
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_image(doc, path, caption=None):
    if not path.exists():
        print(f"  [warn] missing image {path}")
        return
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.space_before = Pt(12)
        cap.paragraph_format.space_after = Pt(4)
        cap.paragraph_format.line_spacing = 1.1
        cap.paragraph_format.keep_with_next = True
        add_runs(cap, caption, size=Pt(9), color=MUTED)
        caption = None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(16)
    p.add_run().add_picture(str(path), width=CONTENT_WIDTH)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.space_after = Pt(14)
        cap.paragraph_format.line_spacing = 1.1
        add_runs(cap, caption, size=Pt(9), color=MUTED)


# --- the parser ---------------------------------------------------------------
IMG = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)\s*$")


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_divider(line):
    cells = split_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


def convert():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")
    doc = Document()
    setup(doc)
    add_page_numbers(doc)

    title_done = False
    pending_caption = None
    i = 0
    n_tables = n_images = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # blank
        if not line.strip():
            i += 1
            continue

        # horizontal rule between sections
        if re.fullmatch(r"-{3,}", line.strip()):
            add_rule(doc)
            i += 1
            continue

        # image
        m = IMG.match(line.strip())
        if m:
            src = (HERE / m.group("src")).resolve()
            add_image(doc, src, pending_caption)
            pending_caption = None
            n_images += 1
            i += 1
            continue

        # table
        if line.lstrip().startswith("|") and i + 1 < len(lines) \
                and is_divider(lines[i + 1]):
            rows = [split_row(line)]
            j = i + 2
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                rows.append(split_row(lines[j]))
                j += 1
            add_table(doc, rows)
            n_tables += 1
            i = j
            continue

        # heading
        hm = re.match(r"^(#{1,4})\s+(.*)$", line)
        if hm:
            level, body = len(hm.group(1)), hm.group(2).strip()
            if body == "Figures":
                brk = doc.add_paragraph()
                brk.add_run().add_break(WD_BREAK.PAGE)
            if level == 1 and not title_done:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(14)
                add_runs(p, f"**{body}**", size=Pt(17))
                title_done = True
            else:
                doc.add_heading("", level=min(level, 3))
                add_runs(doc.paragraphs[-1], body)
            i += 1
            continue

        # figure caption: hold it until the image that follows
        if re.match(r"^\*\*Figure \d+\.\*\*", line.strip()):
            pending_caption = line.strip()
            i += 1
            continue

        # bullet list
        if re.match(r"^\s*[-*]\s+", line):
            indent = len(line) - len(line.lstrip())
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            body = re.sub(r"^\s*[-*]\s+", "", line)
            while i + 1 < len(lines) and lines[i + 1].strip() \
                    and not re.match(r"^\s*[-*]\s+|^#|^\||^!\[|^-{3,}",
                                     lines[i + 1]):
                i += 1
                body += " " + lines[i].strip()
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, body)
            i += 1
            continue

        # paragraph: join wrapped lines
        body = line.strip()
        while i + 1 < len(lines) and lines[i + 1].strip() \
                and not re.match(r"^\s*[-*]\s+|^#|^\||^!\[|^-{3,}|^\*\*Figure",
                                 lines[i + 1]):
            i += 1
            body += " " + lines[i].strip()
        if body.startswith("**This is a modelling study"):
            add_callout(doc, body)
            i += 1
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, body)
        i += 1

    doc.save(OUT)
    size_kb = OUT.stat().st_size / 1024
    print(f"wrote {OUT.name}  ({size_kb:.0f} kB, {n_tables} tables, "
          f"{n_images} figures)")
    return doc


if __name__ == "__main__":
    convert()
