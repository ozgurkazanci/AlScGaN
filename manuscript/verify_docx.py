"""Inspect the generated .docx, since no renderer is available on this machine."""

import re
import pathlib
import zipfile

from docx import Document
from docx.shared import Emu

P = pathlib.Path(__file__).parent / "MANUSCRIPT.docx"
doc = Document(P)

heads, empty_heads, blanks = [], 0, 0
for p in doc.paragraphs:
    st = p.style.name
    if st.startswith("Heading"):
        heads.append((st, p.text.strip()))
        if not p.text.strip():
            empty_heads += 1
    elif not p.text.strip():
        blanks += 1

print(f"paragraphs: {len(doc.paragraphs)}  (blank: {blanks})")
print(f"tables: {len(doc.tables)}")
print(f"headings: {len(heads)}  (empty: {empty_heads})")
print()
print("outline:")
for st, txt in heads:
    lvl = int(st[-1])
    print("   " * (lvl - 1) + f"H{lvl}  {txt[:66]}")

print("\ntables:")
for i, t in enumerate(doc.tables, 1):
    hdr = " | ".join(c.text.strip()[:18] for c in t.rows[0].cells)
    print(f"  {i}. {len(t.rows)}x{len(t.columns)}  [{hdr}]")

with zipfile.ZipFile(P) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    total = sum(z.getinfo(n).file_size for n in media)
print(f"\nembedded images: {len(media)} ({total/1024/1024:.1f} MB)")

sec = doc.sections[0]
print(f"page: {sec.page_width.cm:.1f} x {sec.page_height.cm:.1f} cm, "
      f"margins {sec.left_margin.cm:.1f} cm")

first = next(p for p in doc.paragraphs if p.text.strip())
print(f"title: {first.text.strip()[:78]}")
# a caption is "Figure N." at the start of a paragraph; plain prose like
# "Figure 3 compares ..." is a cross-reference, not a caption
CAP_RE = re.compile(r"^Figure [0-9]+\.")
caps = [p.text.strip() for p in doc.paragraphs
        if CAP_RE.match(p.text.strip())]
print(f"figure captions: {len(caps)}")
for c in caps:
    print(f"  {c[:74]}")

problems = []
if empty_heads:
    problems.append(f"{empty_heads} empty heading paragraphs")
if len(media) != 6:
    problems.append(f"expected 6 images, found {len(media)}")
if len(caps) != 6:
    problems.append(f"expected 6 figure captions, found {len(caps)}")
stray = [p.text for p in doc.paragraphs if "**" in p.text or "![" in p.text]
if stray:
    problems.append(f"{len(stray)} paragraphs with unparsed markdown")
print("\nPROBLEMS:" if problems else "\nno structural problems found")
for x in problems:
    print("  -", x)
